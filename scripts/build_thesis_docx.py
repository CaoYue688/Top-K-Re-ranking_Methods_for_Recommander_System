"""Build the FH-Wedel-style thesis DOCX from the audited manuscript and results.

The source guideline DOCX is copied first and its package-level styles, theme,
numbering, and settings are retained. Only the instructional body is replaced.
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "thesis" / "manuscript.md"
AGG = ROOT / "outputs" / "thesis_pos4_neg2_traincore5_dataseed2026" / "aggregate"
FIGURES = ROOT / "outputs" / "thesis" / "figures"
DEFAULT_TEMPLATE = Path(r"C:\Users\Aroeh\Downloads\Richtlinie 3.0 Word-Vorlage (Stand 25.04.2024).docx")
DEFAULT_OUTPUT = ROOT / "outputs" / "thesis" / "Diversitaetsorientiertes_Re-Ranking_Masterarbeit_Yue_Cao.docx"

HEADER_TEXT = "Diversitätsorientiertes Re-Ranking"
DOC_TITLE = "Diversitätsorientiertes Re-Ranking in Recommender-Systemen"
DOC_SUBJECT = "Masterarbeit – Accuracy–Diversity-Trade-off"
SCHOOL_NAME = "Fachhochschule Wedel"
THESIS_TITLE = "Diversitätsorientiertes Re-Ranking\nin Recommender-Systemen"
THESIS_SUBTITLE = "Eine vergleichende Analyse des Accuracy–Diversity-Trade-offs"
DEGREE_TEXT = "Masterarbeit\nim Studiengang Data Science & Artificial Intelligence"
TITLE_METADATA = [
    ("Vorgelegt von", os.environ.get("THESIS_AUTHOR", "Yue Cao")),
    ("Matrikelnummer", os.environ.get("THESIS_STUDENT_ID", "[Matrikelnummer ergänzen]")),
    ("Anschrift", os.environ.get("THESIS_ADDRESS", "[Anschrift ergänzen]")),
    ("E-Mail", os.environ.get("THESIS_EMAIL", "[E-Mail-Adresse ergänzen]")),
    ("Studiensemester", "5. Fachsemester (Sommersemester 2026)"), ("Erstbetreuung", "Prof. Dr. Gerd Beuster"),
    ("Zweit-/Praxisbetreuung", "–"),
    ("Abgabedatum", "25.08.2026"),
]
BODY_START_HEADING = "1 Einleitung"
BIBLIOGRAPHY_HEADING = "Literaturverzeichnis"
APPENDIX_PREFIX = "Anhang"
DECLARATION_HEADINGS = {"Eidesstattliche Erklärung"}
AUXILIARY_HEADINGS = {"Hilfsmittelverzeichnis"}
FIGURE_LABEL = "Abbildung"
TABLE_LABEL = "Tabelle"
FONT_REGULAR = "Arial"
FONT_EAST_ASIA = "Arial"
LANGUAGE_TAG = "de-DE"
AUTO_HYPHENATION = True
TOC_PLACEHOLDER = "Inhaltsverzeichnis in Word aktualisieren"
FIGURE_TOC_PLACEHOLDER = "Abbildungsverzeichnis in Word aktualisieren"
TABLE_TOC_PLACEHOLDER = "Tabellenverzeichnis in Word aktualisieren"

FOOTNOTES = {
    "[[FN001]]": "Vgl. McNee/Riedl/Konstan (2006), S. 1097–1101.",
    "[[FN002]]": "Vgl. Rendle et al. (2009), S. 452–461.",
    "[[FN003]]": "Vgl. Koren/Bell/Volinsky (2009), S. 30–37.",
    "[[FN004]]": "Vgl. Kaminskas/Bridge (2017), S. 1–42.",
    "[[FN005]]": "Vgl. Steck (2018), S. 154–162.",
    "[[FN006]]": "Vgl. Carbonell/Goldstein (1998), S. 335–336.",
    "[[FN007]]": "Vgl. Santos/Peng/Macdonald/Ounis (2010), S. 87–99.",
    "[[FN008]]": "Vgl. Krichene/Rendle (2020), S. 1748–1757.",
    "[[FN009]]": "Vgl. Abdollahpouri/Burke/Mobasher (2017), S. 42–46.",
    "[[FN010]]": "Vgl. Herlocker et al. (2004), S. 5–53.",
    "[[FN011]]": "Vgl. Ziegler et al. (2005), S. 22–32.",
    "[[FN012]]": "Vgl. Vargas/Castells (2011), S. 109–116.",
    "[[FN013]]": "Vgl. Steck (2018), S. 154–162.",
    "[[FN014]]": "Vgl. Vig/Sen/Riedl (2012), S. 1–44.",
    "[[FN015]]": "Vgl. Zangerle/Bauer (2022), S. 1–38.",
    "[[FN016]]": "Zur Bootstrap-Methodik vgl. Efron/Tibshirani (1993); zur multiplen Testkorrektur vgl. Holm (1979), S. 65–70; zu standardisierten Effektgrößen vgl. Cohen (1988).",
    "[[FN017]]": "Vgl. Harper/Konstan (2015), S. 1–19, sowie die offizielle MovieLens-20M-Datensatzbeschreibung von GroupLens.",
    "[[FN018]]": "Vgl. Jannach/Chen (2026), S. 1–15.",
    "[[FN019]]": "Vgl. Herlocker et al. (2004), S. 5–53, sowie Zangerle/Bauer (2022), S. 1–38.",
    "[[FN020]]": "Vgl. Carbonell/Goldstein (1998), S. 335–336.",
    "[[FN021]]": "Vgl. Kaya/Bridge (2019b), S. 1639–1646.",
    "[[FN022]]": "Vgl. Kaya/Bridge (2019a), S. 151–159.",
    "[[FN023]]": "Vgl. Wang et al. (2023), S. 223–233.",
    "[[FN024]]": "Vgl. Hidasi/Czapp (2023).",
    "[[FN025]]": "Vgl. Kaya/Bridge (2019a), S. 151–159, sowie Carraro/Bridge (2026), S. 1–40.",
    "[[FN026]]": "Vgl. Kaya/Bridge (2019a), S. 151–159.",
    "[[FN027]]": "Vgl. Wang et al. (2023), S. 223–233.",
    "[[FN028]]": "Vgl. Hidasi/Czapp (2023).",
    "[[FN029]]": "Vgl. Carraro/Bridge (2026), S. 1–40.",
    "[[FN030]]": "Vgl. Kunaver/Požrl (2017), S. 154–162.",
    "[[FN031]]": "Vgl. Zhou et al. (2010), S. 4511–4515.",
    "[[FN032]]": "Vgl. Adomavicius/Kwon (2012), S. 896–911.",
    "[[FN033]]": "Vgl. Ge/Delgado-Battenfeld/Jannach (2010), S. 257–260.",
    "[[FN034]]": "Vgl. Singh/Joachims (2018), S. 2219–2228.",
}

TABLE_TITLES = {
    "dataset_stats": "Datensatzstatistik nach zeitlichem Split und trainbasiertem 5-Core",
    "baseline_metrics": "Seedgemittelte Baselinekennzahlen im Hauptversuch",
    "budget_results": "Testresultate der methodenübergreifenden Budgetauswahl",
    "method_comparison_5pct": "ILD-selektierte Testpunkte innerhalb des 5-%-Validierungsbudgets",
    "construct_comparison_5pct": "Konstruktorientierte Testpunkte innerhalb des 5-%-Validierungsbudgets",
    "coverage_comparison_5pct": "Aspektabdeckung bei zielgleicher Validierungsselektion",
    "subgroup_results": "MMR λ=0,40 nach Nutzeraktivität und Profilbreite",
    "segment_lambdas": "Segmentbezogene MMR-Betriebspunkte bei 5 % Accuracy-Budget",
    "robustness_results": "Robustheitsmatrix bei 5-%-Validierungsbudget und Seed 2026",
    "candidate_frontier": "Validierungsselektierte MMR-Punkte nach Kandidatenpoolgröße",
    "tag_results": "Sensitivität des MMR-Trade-offs gegenüber dem Merkmalsraum",
    "runtime_results": "Gemessene Laufzeitprofile des Hauptsweeps",
    "hypotheses": "Ergebnisübersicht der Hypothesenprüfung",
    "protocol_deviations": "Dokumentierte Abweichungen und Einordnung gegenüber dem Exposé",
    "ai_use": "Einsatz generativer KI",
    "parameters": "Parameter des finalen Experimentlaufs",
    "all_budget_methods": "Validierungsselektierte Punkte je Budget, Scope und Zielmetrik",
}


def fmt(value: object, digits: int = 4) -> str:
    if isinstance(value, (float, np.floating)):
        if np.isnan(value):
            return "–"
        return f"{float(value):.{digits}f}".replace(".", ",")
    if isinstance(value, (int, np.integer)):
        return f"{int(value):,}".replace(",", ".")
    return str(value)


def pct(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f} %".replace(".", ",")


def set_cell_text(cell, value: object, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = FONT_REGULAR
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    r.font.size = Pt(7.5 if not bold else 8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for style_name in ("Table Grid", "Tabellenraster"):
        if style_name in doc.styles:
            table.style = style_name
            break
    table.autofit = widths is None
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    mark_header_row(table)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx, value in enumerate(row):
            set_cell_text(cells[cidx], value)
            if ridx % 2:
                shade_cell(cells[cidx], "F3F6F8")
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def mark_header_row(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_field(paragraph, instruction: str, placeholder: str = "Feld in Word aktualisieren") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    result_run = OxmlElement("w:r")
    result_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, result_run, end])


def set_page_number(section, fmt_name: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:pgNumType"))
    if old is not None:
        sect_pr.remove(old)
    num = OxmlElement("w:pgNumType")
    num.set(qn("w:fmt"), fmt_name)
    num.set(qn("w:start"), str(start))
    sect_pr.append(num)


def clear_part(part) -> None:
    for paragraph in list(part.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for table in list(part.tables):
        table._element.getparent().remove(table._element)


def add_page_footer(section, first_page_blank: bool = False) -> None:
    section.footer.is_linked_to_previous = False
    clear_part(section.footer)
    p = section.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE", "1")
    if first_page_blank:
        section.different_first_page_header_footer = True
        section.first_page_footer.is_linked_to_previous = False
        clear_part(section.first_page_footer)


def configure_section(section, left: float = 3.5, right: float = 4.0, top: float = 3.0, bottom: float = 3.0) -> None:
    section.different_first_page_header_footer = False
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def set_header(section, text: str = "") -> None:
    section.header.is_linked_to_previous = False
    clear_part(section.header)
    if text:
        p = section.header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text)
        r.font.name = FONT_REGULAR
        r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(90, 90, 90)


def ensure_style(doc: Document, name: str, base: str | None = None):
    if name in doc.styles:
        style = doc.styles[name]
    else:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base and base in doc.styles:
        style.base_style = doc.styles[base]
    return style


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_REGULAR
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    body = ensure_style(doc, "Thesis Body", "Normal")
    body.font.name = FONT_REGULAR
    body._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    body.font.size = Pt(11)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.widow_control = True

    for name, size, color in [
        ("Heading 1", 16, "1F4E79"), ("Heading 2", 13, "1F4E79"), ("Heading 3", 11, "365F91")
    ]:
        style = doc.styles[name]
        style.font.name = FONT_REGULAR
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(7 if name == "Heading 1" else 5)
    doc.styles["Heading 1"].paragraph_format.page_break_before = True

    caption = ensure_style(doc, "Caption", "Normal")
    caption.font.name = FONT_REGULAR
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    bib = ensure_style(doc, "Bibliography Entry", "Normal")
    bib.font.name = FONT_REGULAR
    bib._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    bib.font.size = Pt(10)
    bib.paragraph_format.left_indent = Cm(0.6)
    bib.paragraph_format.first_line_indent = Cm(-0.6)
    bib.paragraph_format.line_spacing = 1.15
    bib.paragraph_format.space_after = Pt(5)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_caption(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{label} ")
    add_field(p, f"SEQ {label} \\* ARABIC", "1")
    p.add_run(f": {text}")


def add_literal_caption(doc: Document, label: str, number: str, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{label} {number}: {text}")


def add_numbered_equation(doc: Document, number: str, formula: str) -> None:
    p = doc.add_paragraph(style="Thesis Body")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(Cm(6.75), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(13.5), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run(f"\t{formula}\t({number})")


def make_tables(data: pd.DataFrame) -> dict[str, tuple[list[str], list[list[object]]]]:
    primary = data[(data.experiment == "primary_n100_k10_genre") & (data.split == "test") & (data.method == "mmr")]
    baseline = primary[primary["lambda"] == 0]
    bmean = baseline.mean(numeric_only=True)

    dataset_stats = [
        ["Rohbewertungen", fmt(20_000_263)], ["Positive Rohinteraktionen (≥4)", fmt(9_995_410)],
        ["Nutzende nach Core", fmt(134_703)], ["Filme nach Core", fmt(11_851)],
        ["Positive Interaktionen nach Core", fmt(9_952_928)], ["Training", fmt(7_908_519)],
        ["Validierung", fmt(939_551)], ["Test", fmt(1_104_858)],
        ["Explizite Negativinteraktionen", fmt(2_563_528)], ["Neutrale Interaktionen", fmt(7_260_442)],
        ["Genremerkmale", "20"], ["Core-Basis", "chronologische Trainingspartition"],
    ]

    baseline_rows = [
        ["Recall@10", fmt(bmean["recall@10"])], ["NDCG@10", fmt(bmean["ndcg@10"])],
        ["MRR@10", fmt(bmean["mrr@10"])], ["ILD@10", fmt(bmean["ild@10"])],
        ["Kalibrierungsähnlichkeit@10", fmt(bmean["calibration@10"])],
        ["Subtopic Recall@10", fmt(bmean["subtopic_recall@10"])],
        ["Katalogcoverage@10", fmt(bmean["catalog_coverage@10"])],
        ["Exposure-Gini@10", fmt(bmean["exposure_gini@10"])],
        ["Long-Tail-Anteil@10", fmt(bmean["long_tail_share@10"])],
        ["Candidate Recall@100", fmt(bmean["candidate_recall@100"])],
    ]

    budgets = pd.read_csv(AGG / "test_budget_results.csv")
    across = budgets[budgets.scope == "primary_ild_across_methods"].copy()
    budget_rows = []
    for _, r in across.iterrows():
        base_ndcg = r["test_ndcg@10_mean"] - r["delta_ndcg@10_mean"]
        base_ild = r["test_ild@10_mean"] - r["delta_ild@10_mean"]
        budget_rows.append([
            pct(100 * r["budget"], 0), r["method"].upper(), fmt(r["lambda"], 2),
            fmt(r["test_ndcg@10_mean"], 6), pct(100 * r["delta_ndcg@10_mean"] / base_ndcg),
            fmt(r["test_ild@10_mean"], 6), pct(100 * r["delta_ild@10_mean"] / base_ild),
        ])

    method_rows = []
    for _, r in budgets[(budgets.scope == "within_method_ild") & np.isclose(budgets.budget, 0.05)].iterrows():
        method_rows.append([
            {"mmr": "MMR", "xquad": "xQuAD", "calibration": "Kalibrierung"}[r["method"]],
            fmt(r["lambda"], 2), fmt(r["test_ndcg@10_mean"], 6), fmt(r["test_ild@10_mean"], 6),
            fmt(r["test_calibration@10_mean"], 6), fmt(r["test_catalog_coverage@10_mean"], 6),
            fmt(r["test_long_tail_share@10_mean"], 6),
        ])

    target_labels = {
        "ild@10": "ILD@10",
        "subtopic_recall@10": "Subtopic Recall@10",
        "calibration@10": "Kalibrierung@10",
    }
    construct_rows = []
    construct = budgets[
        (budgets.scope == "construct_aligned") & np.isclose(budgets.budget, 0.05)
    ]
    for _, r in construct.iterrows():
        target = str(r["target_metric"])
        test_value = {
            "ild@10": r["test_ild@10_mean"],
            "subtopic_recall@10": r["test_subtopic_recall@10_mean"],
            "calibration@10": r["test_calibration@10_mean"],
        }[target]
        baseline_value = {
            "ild@10": r["baseline_ild@10_mean"],
            "subtopic_recall@10": r["baseline_subtopic_recall@10_mean"],
            "calibration@10": r["baseline_calibration@10_mean"],
        }[target]
        construct_rows.append([
            {"mmr": "MMR", "xquad": "xQuAD", "calibration": "Kalibrierung"}[r["method"]],
            target_labels[target], fmt(r["lambda"], 2), fmt(r["test_ndcg@10_mean"], 6),
            pct(100 * (r["test_ndcg@10_mean"] / r["baseline_ndcg@10_mean"] - 1)),
            fmt(test_value, 6), pct(100 * (test_value / baseline_value - 1)),
        ])

    coverage_data = pd.read_csv(AGG / "coverage_comparison_5pct.csv")
    coverage_rows = []
    for _, r in coverage_data.iterrows():
        coverage_rows.append([
            {"mmr": "MMR", "xquad": "xQuAD", "calibration": "Kalibrierung"}[r["method"]],
            fmt(r["lambda"], 2), fmt(r["test_metric"], 6),
            pct(100 * r["test_metric_delta_relative"]),
            pct(100 * r["test_ndcg_delta_relative"]),
        ])

    chosen = primary[np.isclose(primary["lambda"], 0.4, atol=1e-7)]
    groups = [
        ("Geringe Aktivität", "low_activity"), ("Mittlere Aktivität", "medium_activity"),
        ("Hohe Aktivität", "high_activity"), ("Fokussiertes Profil", "focused_profile"),
        ("Mittleres Profil", "medium_profile"), ("Breites Profil", "broad_profile"),
    ]
    subgroup_rows = []
    for label, key in groups:
        users = int(chosen[f"{key}_users"].mean())
        base_ndcg = baseline[f"{key}_ndcg@10"].mean()
        sel_ndcg = chosen[f"{key}_ndcg@10"].mean()
        base_ild = baseline[f"{key}_ild@10"].mean()
        sel_ild = chosen[f"{key}_ild@10"].mean()
        subgroup_rows.append([
            label, fmt(users), pct(100 * (sel_ndcg / base_ndcg - 1)), pct(100 * (sel_ild / base_ild - 1)),
        ])

    segment_data = pd.read_csv(AGG / "segment_budget_selections.csv")
    segment_labels = {
        "focused_profile": "Fokussiertes Profil",
        "medium_profile": "Mittleres Profil",
        "broad_profile": "Breites Profil",
    }
    segment_rows = []
    for _, r in segment_data[np.isclose(segment_data.budget, 0.05)].iterrows():
        segment_rows.append([
            segment_labels[r["segment"]], fmt(r["lambda"], 2),
            pct(100 * r["validation_ndcg_delta_relative"]),
            pct(100 * r["validation_ild_delta_relative"]),
        ])

    def select_seed(experiment: str, method: str) -> tuple[pd.Series, pd.Series]:
        val = data[(data.experiment == experiment) & (data.split == "val") & (data.method == method) & (data.seed == 2026)]
        k = int(val.top_k.iloc[0])
        b = val[val["lambda"] == 0].iloc[0]
        feasible = val[val[f"ndcg@{k}"] >= 0.95 * b[f"ndcg@{k}"]]
        sel = feasible.sort_values([f"ild@{k}", f"ndcg@{k}"], ascending=False).iloc[0]
        test = data[(data.experiment == experiment) & (data.split == "test") & (data.method == method) & (data.seed == 2026)]
        t = test[np.isclose(test["lambda"], sel["lambda"], atol=1e-7)].iloc[0]
        tb = test[test["lambda"] == 0].iloc[0]
        return t, tb

    robustness_rows = []
    for exp, label in [
        ("robust_n50_k10_genre", "N=50, K=10"), ("primary_n100_k10_genre", "N=100, K=10"),
        ("robust_n200_k10_genre", "N=200, K=10"), ("robust_n100_k5_genre", "N=100, K=5"),
        ("robust_n100_k20_genre", "N=100, K=20"),
    ]:
        t, tb = select_seed(exp, "mmr")
        k = int(t.top_k)
        robustness_rows.append([
            label, fmt(t["lambda"], 2), pct(100 * (t[f"ndcg@{k}"] / tb[f"ndcg@{k}"] - 1)),
            pct(100 * (t[f"ild@{k}"] / tb[f"ild@{k}"] - 1)), fmt(tb[f"candidate_recall@{int(tb.candidate_k)}"], 4),
        ])

    frontier_data = pd.read_csv(AGG / "candidate_pool_budget_frontier.csv")
    frontier_rows = []
    for _, r in frontier_data[np.isclose(frontier_data.budget, 0.05)].iterrows():
        frontier_rows.append([
            fmt(int(r["candidate_k"])), fmt(r["lambda"], 2),
            fmt(r["validation_ndcg"], 6), fmt(r["validation_ild"], 6),
            "ja" if r["outward_vs_previous_pool"] == 1 else "nein",
        ])

    tag_rows = []
    for exp, label in [
        ("tag_sensitivity_n100_k10_genre", "Genre"),
        ("tag_sensitivity_n100_k10_tag_genome_svd64", "Tag Genome SVD64"),
    ]:
        val = data[(data.experiment == exp) & (data.split == "val") & (data.method == "mmr") & (data.seed == 2026)]
        vb = val[val["lambda"] == 0].iloc[0]
        feasible = val[val["ndcg@10"] >= 0.95 * vb["ndcg@10"]]
        sel = feasible.sort_values(["feature_ild@10", "ndcg@10"], ascending=False).iloc[0]
        test = data[(data.experiment == exp) & (data.split == "test") & (data.method == "mmr") & (data.seed == 2026)]
        t = test[np.isclose(test["lambda"], sel["lambda"], atol=1e-7)].iloc[0]
        tb = test[test["lambda"] == 0].iloc[0]
        tag_rows.append([
            label, fmt(sel["lambda"], 2), pct(100 * (t["ndcg@10"] / tb["ndcg@10"] - 1)),
            fmt(t["feature_ild@10"], 6), pct(100 * (t["feature_ild@10"] / tb["feature_ild@10"] - 1)),
        ])

    hypothesis_rows = [
        ["H1", "Alle drei Kernverfahren erhöhen mindestens eine Diversity-Metrik.", "unterstützt", "Konstruktorientierte 5-%-Punkte liegen jeweils über der Baseline; NDCG sinkt."],
        ["H2", "xQuAD erzielt stärkere Genre-Coverage-Gewinne als MMR.", "unterstützt", "Subtopic Recall: xQuAD +9,45 %, MMR +7,17 %; Kalibrierung erreicht explorativ +9,79 %."],
        ["H3", "Kalibriertes Re-Ranking erhält individuelle Interessenanteile besser.", "unterstützt", "Kalibrierungsähnlichkeit +12,35 % gegenüber +2,68 % bei global ILD-selektiertem MMR."],
        ["H4", "Größere Kandidatenpools verschieben die Pareto-Front nach außen.", "nicht unterstützt", "N=100 verbessert N=50 bei allen Budgets; N=200 verbessert N=100 nur beim 1-%-Punkt."],
        ["H5", "Ein globales λ ist nicht für alle Profilsegmente optimal.", "unterstützt", "5-%-Selektion: λ=0,35 / 0,40 / 0,45 für fokussierte / mittlere / breite Profile."],
    ]

    runtime_rows = []
    runtime_source = data[
        (data.experiment == "primary_n100_k10_genre")
        & (data.split == "test")
        & np.isclose(data["lambda"], 0.0)
    ]
    for method, frame in runtime_source.groupby("method"):
        runtime_rows.append([
            {"mmr": "MMR", "xquad": "xQuAD", "calibration": "Kalibrierung"}[method],
            fmt(frame["rerank_seconds_all_lambdas"].mean(), 2),
            fmt(frame["amortized_rerank_ms_per_user_config"].mean(), 4),
            fmt(frame["peak_traced_memory_mb"].mean(), 1),
        ])

    protocol_rows = [
        ["Hypothesen", "Die überarbeitete Exposé-Fassung vom 20. August 2026 entstand nach wesentlichen Implementierungs- und Analysearbeiten.", "H1–H5 als hypothesengeleiteter, nicht präregistrierter Rahmen; resultatsnahe Aussagen explorativ.", "keine konfirmatorische oder zeitlich vorgezogene Evidenz beansprucht; H4 nicht unterstützt"],
        ["Core-Filter", "Frühe Pipeline nutzte den vollständigen positiven Verlauf.", "Finaler 5-Core ausschließlich auf chronologischem Training.", "verhindert Look-ahead-Leakage"],
        ["λ-Konvention", "Exposé: λ gewichtet Relevanz; Implementierung: λ gewichtet Diversität.", "Konvention der Implementierung vollständig dokumentiert; λ=0 ist Baseline.", "Grid deckt dieselben Mischungen ab, Zahlen sind nicht direkt gleich benannt"],
        ["Aspektmodell", "xQuAD war als aspektorientierte Methode vorgesehen.", "Binäre Genre-Hard-Coverage statt probabilistischer Soft-Coverage.", "Interpretation auf erste Genreabdeckung begrenzt"],
        ["BPR-Tuning", "Exposé sah Sensitivitätsprüfungen vor.", "BPR als fester Referenz-Retriever; kein Modellvergleich behauptet.", "absolute Retrievalqualität bleibt Limitation"],
        ["Sanity Baseline", "Zufälliges Re-Ranking war optional vorgesehen.", "Nicht in die Kernanalyse aufgenommen; λ=0 und Diversity-only dienen als technische Grenzpunkte.", "keine Aussage über Überlegenheit gegenüber Zufallslisten"],
    ]

    ai_rows = [[
        "Codex Desktop (OpenAI)",
        "Webdienst, Stand 08/2026",
        "Übertragung eigenständig auf Chinesisch formulierter Gedanken ins Deutsche; sprachliche Überarbeitung; Entwicklung von Suchbegriffen und strukturierte Erschließung von Fachpublikationen; methodisches und statistisches Sparring; Codevorschläge, Fehlersuche, Tests sowie Prüfung des Manuskripts. Alle Ergebnisse wurden eigenständig kontrolliert, überarbeitet oder verworfen.",
        "Zusammenfassung und Kap. 1–8; Implementierung zu Kap. 4–5; Anhänge und Hilfsmittelverzeichnis.",
    ]]

    parameter_rows = [
        ["Datensatz", "MovieLens 20M"], ["Feedback", "positiv ≥4; negativ ≤2; neutral 2,5–3"],
        ["Core", "iterativer 5-Core auf chronologischem Training"], ["Seeds", "2026, 2027, 2028"],
        ["Basismodell", "BPR-MF"], ["Latente Dimension", "64"], ["Epochen", "10"],
        ["Batchgröße", "8.192"], ["Explizite Negativquote", "0,5"], ["Gerät", "CUDA / RTX 3070"],
        ["Hauptkonfiguration", "N=100, K=10"], ["Methoden", "MMR, xQuAD, Kalibrierung"],
        ["λ-Hauptraster", "0,00 bis 1,00 in 0,05"], ["Accuracy-Budgets", "1 %, 3 %, 5 %, 10 %"],
        ["Bootstrap", "200 Replikationen Hauptkurven; 100 Robustheitskurven"], ["Primäre Endpunkte", "NDCG@10, ILD@10"],
        ["Vorzeichentest", "zweiseitige Normalapproximation mit Stetigkeitskorrektur"],
        ["Holm-Familie", "vier primäre ILD-Budgetpunkte; sekundäre Tests explorativ"],
    ]

    all_budget_rows = []
    for _, r in budgets.sort_values(["budget", "scope", "method"]).iterrows():
        scope_label = {
            "primary_ild_across_methods": "primär über Methoden",
            "within_method_ild": "ILD je Methode",
            "construct_aligned": "Konstrukt je Methode",
        }[r["scope"]]
        all_budget_rows.append([
            pct(100 * r["budget"], 0), scope_label,
            {"mmr": "MMR", "xquad": "xQuAD", "calibration": "Kal."}[r["method"]], fmt(r["lambda"], 2),
            target_labels[str(r["target_metric"])], fmt(r["test_ndcg@10_mean"], 6), fmt(r["test_ild@10_mean"], 6),
        ])

    return {
        "dataset_stats": (["Kennzahl", "Wert"], dataset_stats),
        "baseline_metrics": (["Metrik", "Mittelwert"], baseline_rows),
        "budget_results": (["Budget", "Methode", "λ", "NDCG", "Δ NDCG rel.", "ILD", "Δ ILD rel."], budget_rows),
        "method_comparison_5pct": (["Methode", "λ", "NDCG", "ILD", "Kalibrierung", "Coverage", "Long Tail"], method_rows),
        "construct_comparison_5pct": (["Methode", "Ziel", "λ", "NDCG", "Δ NDCG", "Zielwert", "Δ Zielwert"], construct_rows),
        "coverage_comparison_5pct": (["Methode", "λ", "Subtopic Recall", "Δ Subtopic", "Δ NDCG"], coverage_rows),
        "subgroup_results": (["Gruppe", "N", "Δ NDCG rel.", "Δ ILD rel."], subgroup_rows),
        "segment_lambdas": (["Profilsegment", "λ", "Δ NDCG Val.", "Δ ILD Val."], segment_rows),
        "robustness_results": (["Einstellung", "λ", "Δ NDCG rel.", "Δ ILD rel.", "Candidate Recall"], robustness_rows),
        "candidate_frontier": (["N", "λ", "Val.-NDCG", "Val.-ILD", "ILD höher als vorheriges N"], frontier_rows),
        "tag_results": (["Merkmalsraum", "λ", "Δ NDCG rel.", "Feature-ILD", "Δ Feature-ILD rel."], tag_rows),
        "runtime_results": (["Methode", "Sekunden für Sweep", "ms/Nutzer/Konfiguration", "Peak MB"], runtime_rows),
        "hypotheses": (["Hyp.", "Aussage", "Entscheidung", "Evidenz"], hypothesis_rows),
        "protocol_deviations": (["Bereich", "Ausgangslage", "Finale Behandlung", "Auswirkung"], protocol_rows),
        "ai_use": (["Werkzeug (Anbieter)", "Version / Stand", "Zweck", "Stelle in der Arbeit"], ai_rows),
        "parameters": (["Parameter", "Wert"], parameter_rows),
        "all_budget_methods": (["Budget", "Scope", "Methode", "λ", "Ziel", "Test-NDCG", "Test-ILD"], all_budget_rows),
    }


def add_markdown_table(doc: Document, lines: list[str], start: int) -> int:
    block = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        block.append(lines[i].strip())
        i += 1
    parsed = [[c.strip() for c in row.strip("|").split("|")] for row in block]
    if len(parsed) >= 2 and all(set(c) <= {"-", ":"} for c in parsed[1]):
        headers = parsed[0]
        rows = parsed[2:]
        add_table(doc, headers, rows)
    return i


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Cm(1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SCHOOL_NAME)
    r.bold = True
    r.font.name = FONT_REGULAR
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Cm(2.6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(THESIS_TITLE)
    r.bold = True
    r.font.name = FONT_REGULAR
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Cm(0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(THESIS_SUBTITLE)
    r.font.name = FONT_REGULAR
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Cm(1.6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(DEGREE_TEXT)
    r.font.name = FONT_REGULAR
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    r.font.size = Pt(12)

    meta = TITLE_METADATA
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row, (key, value) in zip(table.rows, meta):
        row.cells[0].width = Cm(5.2)
        row.cells[1].width = Cm(7.2)
        set_cell_text(row.cells[0], key, bold=True)
        set_cell_text(row.cells[1], value)
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "nil")
                borders.append(el)
    mark_header_row(table)


def build(template: Path, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)
    doc = Document(output)
    clear_body(doc)
    configure_styles(doc)
    doc.core_properties.title = DOC_TITLE
    doc.core_properties.subject = DOC_SUBJECT
    doc.core_properties.author = "Yue Cao"
    doc.core_properties.keywords = "Recommender-Systeme, Diversität, MMR, xQuAD, Kalibrierung"

    title_section = doc.sections[0]
    configure_section(title_section)
    set_header(title_section)
    title_section.footer.is_linked_to_previous = False
    clear_part(title_section.footer)
    title_section.different_first_page_header_footer = True
    add_title_page(doc)

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front)
    set_page_number(front, "lowerRoman", 1)
    add_page_footer(front)
    set_header(front, HEADER_TEXT)

    raw = MANUSCRIPT.read_text(encoding="utf-8")
    first_break = raw.index("@@PAGEBREAK@@")
    content = raw[first_break + len("@@PAGEBREAK@@"):]
    content = content.replace(
        f"@@PAGEBREAK@@\n\n# {BODY_START_HEADING}",
        f"# {BODY_START_HEADING}",
    )
    lines = content.splitlines()
    data = pd.read_csv(AGG / "all_thesis_results.csv")
    tables = make_tables(data)
    body_started = False
    in_bibliography = False
    in_appendix = False
    appendix_letter = ""
    appendix_table_numbers: dict[str, int] = {}
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "@@PAGEBREAK@@":
            doc.add_page_break()
            i += 1
            continue
        if stripped == "@@TOC@@":
            p = doc.add_paragraph()
            add_field(p, 'TOC \\o "1-3" \\h \\z \\u', TOC_PLACEHOLDER)
            i += 1
            continue
        if stripped == "@@TOC_FIGURES@@":
            p = doc.add_paragraph()
            add_field(p, f'TOC \\h \\z \\c "{FIGURE_LABEL}"', FIGURE_TOC_PLACEHOLDER)
            i += 1
            continue
        if stripped == "@@TOC_TABLES@@":
            p = doc.add_paragraph()
            add_field(p, f'TOC \\h \\z \\c "{TABLE_LABEL}"', TABLE_TOC_PLACEHOLDER)
            i += 1
            continue
        fig_match = re.fullmatch(r"@@FIG:([^|]+)\|(.+)@@", stripped)
        if fig_match:
            filename, caption = fig_match.groups()
            pic = doc.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic.add_run()
            shape = run.add_picture(str(FIGURES / filename), width=Cm(13.1))
            clean = re.sub(rf"^{re.escape(FIGURE_LABEL)}\s+\d+:\s*", "", caption)
            shape._inline.docPr.set("title", clean)
            shape._inline.docPr.set("descr", clean)
            add_caption(doc, FIGURE_LABEL, clean)
            i += 1
            continue
        table_match = re.fullmatch(r"@@TABLE:([^@]+)@@", stripped)
        if table_match:
            name = table_match.group(1)
            if in_appendix:
                appendix_table_numbers[appendix_letter] = appendix_table_numbers.get(appendix_letter, 0) + 1
                add_literal_caption(
                    doc,
                    TABLE_LABEL,
                    f"{appendix_letter}.{appendix_table_numbers[appendix_letter]}",
                    TABLE_TITLES[name],
                )
            elif name != "ai_use":
                add_caption(doc, TABLE_LABEL, TABLE_TITLES[name])
            headers, rows = tables[name]
            add_table(doc, headers, rows, widths=[2.6, 2.0, 5.1, 3.8] if name == "ai_use" else None)
            i += 1
            continue
        equation_match = re.fullmatch(r"@@EQ:(\d+)\|(.+)@@", stripped)
        if equation_match:
            number, formula = equation_match.groups()
            add_numbered_equation(doc, number, formula)
            i += 1
            continue
        if stripped.startswith("|"):
            i = add_markdown_table(doc, lines, i)
            continue
        if stripped.startswith("# "):
            heading = stripped[2:]
            if heading == BODY_START_HEADING and not body_started:
                body = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(body)
                set_page_number(body, "decimal", 1)
                add_page_footer(body)
                set_header(body, HEADER_TEXT)
                body_started = True
            p = doc.add_paragraph(heading, style="Heading 1")
            if heading == BIBLIOGRAPHY_HEADING:
                in_bibliography = True
                in_appendix = False
            elif heading in AUXILIARY_HEADINGS:
                in_bibliography = False
                in_appendix = False
            elif heading.startswith(APPENDIX_PREFIX) or heading in DECLARATION_HEADINGS:
                in_bibliography = False
                appendix_match = re.match(rf"{re.escape(APPENDIX_PREFIX)}\s+([A-Z])", heading)
                if appendix_match:
                    in_appendix = True
                    appendix_letter = appendix_match.group(1)
                elif heading in DECLARATION_HEADINGS:
                    in_appendix = False
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 2")
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 3")
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:])
            i += 1
            continue
        style = "Bibliography Entry" if in_bibliography else "Thesis Body"
        p = doc.add_paragraph(style=style)
        p.add_run(stripped)
        i += 1

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    # Enable automatic hyphenation and embed a compatibility-safe language tag.
    auto_hyphen = settings.find(qn("w:autoHyphenation"))
    if auto_hyphen is None:
        auto_hyphen = OxmlElement("w:autoHyphenation")
        settings.append(auto_hyphen)
    auto_hyphen.set(qn("w:val"), "true" if AUTO_HYPHENATION else "false")
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            r_pr = run._r.get_or_add_rPr()
            lang = r_pr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                r_pr.append(lang)
            lang.set(qn("w:val"), LANGUAGE_TAG)

    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--write-footnote-map", type=Path)
    args = parser.parse_args()
    build(args.template, args.output)
    if args.write_footnote_map:
        args.write_footnote_map.parent.mkdir(parents=True, exist_ok=True)
        args.write_footnote_map.write_text(
            "\n".join(f"{marker}\t{text}" for marker, text in FOOTNOTES.items()) + "\n",
            encoding="utf-8",
        )
    print(f"[OK] wrote {args.output}")


if __name__ == "__main__":
    main()
