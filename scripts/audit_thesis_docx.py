"""Fail-fast structural audit for the generated thesis DOCX."""

from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS = {"w": W, "wp": WP}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()
    failures: list[str] = []
    with zipfile.ZipFile(args.docx) as archive:
        if archive.testzip() is not None:
            failures.append("ZIP CRC failure")
        document_xml = archive.read("word/document.xml")
        root = etree.fromstring(document_xml)
        xml_text = document_xml.decode("utf-8")
        for marker in ("[[FN", "@@FIG", "@@TABLE", "@@TOC", "@@PAGEBREAK"):
            if marker in xml_text:
                failures.append(f"unresolved marker: {marker}")
        refs = root.xpath(".//w:footnoteReference", namespaces=NS)
        if len(refs) != 29:
            failures.append(f"expected 29 footnote references, got {len(refs)}")
        foot_root = etree.fromstring(archive.read("word/footnotes.xml"))
        notes = [
            n for n in foot_root.xpath(".//w:footnote", namespaces=NS)
            if int(n.get(f"{{{W}}}id")) >= 1
        ]
        if len(notes) != 29:
            failures.append(f"expected 29 defined footnotes, got {len(notes)}")
        images = root.xpath(".//wp:docPr", namespaces=NS)
        if len(images) != 7:
            failures.append(f"expected 7 images, got {len(images)}")
        if any(not (img.get("descr") or "").strip() for img in images):
            failures.append("image without alt text")
        tables = root.xpath(".//w:tbl", namespaces=NS)
        if len(tables) != 20:
            failures.append(f"expected 20 tables, got {len(tables)}")
        if any(not tbl.xpath("./w:tr[1]/w:trPr/w:tblHeader", namespaces=NS) for tbl in tables):
            failures.append("table without marked header row")
        sections = root.xpath(".//w:sectPr", namespaces=NS)
        formats = [
            (s.find("w:pgNumType", namespaces=NS).get(f"{{{W}}}fmt") if s.find("w:pgNumType", namespaces=NS) is not None else None)
            for s in sections
        ]
        if len(sections) != 3 or formats[-2:] != ["lowerRoman", "decimal"]:
            failures.append(f"unexpected page-number sections: {formats}")

    doc = Document(args.docx)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    all_text = body_text + "\n" + table_text
    word_count = len(re.findall(r"\b[\wÄÖÜäöüß]+\b", all_text))
    headings = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    if word_count < 14_000:
        failures.append(f"unexpectedly short document: {word_count} words")
    if headings < 90:
        failures.append(f"unexpected heading count: {headings}")
    for phrase in ("Richtlinie 3.0", "Formale Gestaltungshinweise", "Beispiele für schriftliche Ausarbeitungen"):
        if phrase in all_text:
            failures.append(f"template instruction leaked: {phrase}")
    for required in (
        "16,30 %", "3,80 %", "134.703", "9.952.928", "Erklärung zur Nutzung generativer KI",
        "[Matrikelnummer ergänzen]", "[Erstbetreuung ergänzen]",
    ):
        if required not in all_text:
            failures.append(f"required text missing: {required}")
    if failures:
        raise SystemExit("AUDIT FAILED\n- " + "\n- ".join(failures))
    print(
        "AUDIT OK | "
        f"words={word_count} paragraphs={len(doc.paragraphs)} headings={headings} "
        f"tables={len(doc.tables)} images=7 sections={len(doc.sections)} footnotes=29 "
        f"bytes={args.docx.stat().st_size}"
    )


if __name__ == "__main__":
    main()
